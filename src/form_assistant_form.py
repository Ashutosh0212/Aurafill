"""Form filling assistant with document editing and chat support.

This module implements a Streamlit interface that combines form filling,
document editing, and chat assistance using vector database and LLM support.
"""

import streamlit as st
from pathlib import Path
import tempfile
from typing import List, Dict, Any, Tuple, Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor
import json
import os
from docx import Document as DocxDocument
import re
from io import BytesIO
import logging
import shutil
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import time
import traceback
import pandas as pd
import atexit

# Import utilities
from src.utils.embeddings import EmbeddingManager
from src.utils.db_utils import get_vectordb, get_available_databases
from src.utils.llm_utils import LLMManager
from src.config.config import QUERY_MODELS
from src.config.users import authenticate
from src.utils.user_logs import UserLogger

# Constants
SUPPORTED_FORMATS = ["docx"]
SECTION_PATTERN = r'^\d+\.\s+.*?(?:\s*\(\d+\s*words\):)?$'

# Global variables
_TEMPLATE_PATH = None

# Update logging configuration
logging.basicConfig(
    filename='logs/app.log',  # Change log file location
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Update file paths to use absolute paths
BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
LOGS_DIR = BASE_DIR / "logs"

# Create necessary directories
DATA_DIR.mkdir(exist_ok=True)
LOGS_DIR.mkdir(exist_ok=True)

# Update temporary file handling
def get_temp_file_path(suffix: str) -> str:
    """Get temporary file path in data directory."""
    return str(DATA_DIR / f"temp_{int(time.time())}{suffix}")

def cleanup_files():
    """Clean up temporary files safely"""
    try:
        global _TEMPLATE_PATH
        if _TEMPLATE_PATH and os.path.exists(_TEMPLATE_PATH):
            os.unlink(_TEMPLATE_PATH)
            _TEMPLATE_PATH = None
    except Exception as e:
        logging.error(f"Cleanup error: {str(e)}")

# Register cleanup
atexit.register(cleanup_files)

def init_session_state():
    """Initialize session state variables."""
    defaults = {
        "messages": [],
        "form_data": {},
        "document_content": "",
        "vectordb": None,
        "chain": None,
        "processing": False,
        "authenticated": False,
        "username": None,
        "user_role": None,
        "template_path": None,
        "document_analysis": None,
        "sections": [],
        "current_section": None,
        "initialized": True
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

# Initialize session state
init_session_state()

# Initialize managers
embedding_manager = EmbeddingManager()
llm_manager = LLMManager()

async def process_query_async(query: str) -> Dict[str, Any]:
    """Process query asynchronously."""
    if not st.session_state.chain:
        return {"answer": "Please load a database first!", "extracted_content": [], "source_count": 0}
    
    try:
        result = await llm_manager.process_query(st.session_state.chain, query)
        return result
    except Exception as e:
        return {
            "answer": f"Error processing query: {str(e)}",
            "extracted_content": [],
            "source_count": 0
        }

def update_form_field(field_name: str, value: str):
    """Update form field in session state."""
    if "form_data" not in st.session_state:
        st.session_state.form_data = {}
    st.session_state.form_data[field_name] = value

def save_document():
    """Save the document content."""
    if st.session_state.document_content:
        try:
            with open("document.txt", "w", encoding="utf-8") as f:
                f.write(st.session_state.document_content)
            st.success("Document saved successfully!")
        except Exception as e:
            st.error(f"Error saving document: {str(e)}")

# Login section
if not st.session_state.authenticated:
    col1, col2 = st.columns(2)
    
    with col1:
        st.title("AI Form Assistant made by ashu")
        username = st.text_input("Username", placeholder="Enter your username")
        password = st.text_input("Password", type="password", placeholder="Enter your password")
        if st.button("Login", use_container_width=True):
            is_authenticated, role = authenticate(username, password)
            if is_authenticated:
                st.session_state.authenticated = True
                st.session_state.username = username
                st.session_state.user_role = role
                st.success("Login successful!")
                st.rerun()
            else:
                st.error("Invalid credentials")
    
    with col2:
        st.markdown("""
            <div style='padding: 2rem; color: #e2e8f0;'>
                <h2>Form Assistant</h2>
                <p>An AI-powered assistant to help you fill forms and edit documents. made by ashu</p>
                <ul>
                    <li>Intelligent form filling</li>
                    <li>Document editing capabilities, and regenerating capabilities in it</li>
                    <li>Chat assistance</li>
                    <li>Vector database integration</li>
                </ul>
            </div>
        """, unsafe_allow_html=True)
    
    st.stop()

# Add CSS for better window fitting
st.markdown("""
<style>
    .main .block-container {
        padding: 1rem;
        max-width: 100%;
    }
    
    .column-container {
        height: calc(100vh - 100px);
        overflow-y: auto;
        padding: 1rem;
        background-color: #1e293b;
        border-radius: 8px;
        margin: 0.5rem;
    }
    
    .document-preview {
        height: calc(100vh - 150px);
        overflow-y: auto;
        background-color: #ffffff;
        color: #000000;
        padding: 2rem;
        border-radius: 4px;
        font-family: 'Times New Roman', serif;
    }
    
    .chat-container {
        height: calc(100vh - 150px);
        display: flex;
        flex-direction: column;
    }
    
    .chat-messages {
        flex-grow: 1;
        overflow-y: auto;
        margin-bottom: 1rem;
    }
    
    .stButton button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

def create_database(db_name: str, files: List[tempfile._TemporaryFileWrapper]) -> None:
    """Create a new vector database from multiple files."""
    try:
        texts = []
        progress_bar = st.progress(0, text="Processing files...")
        
        # Process files
        for idx, file in enumerate(files):
            try:
                # Create a temporary file for each uploaded file
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.name).suffix) as tmp_file:
                    tmp_file.write(file.getvalue())
                    tmp_path = tmp_file.name

                if file.name.endswith('.docx'):
                    # Open the temporary file with python-docx
                    doc = DocxDocument(tmp_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    texts.append(text)
                elif file.name.endswith('.txt'):
                    with open(tmp_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    texts.append(text)

                # Clean up temporary file
                os.unlink(tmp_path)
                
                # Update progress
                progress = (idx + 1) / len(files)
                progress_bar.progress(progress, text=f"Still Processing files... ({idx + 1}/{len(files)})")

            except Exception as e:
                st.error(f"Error processing file {file.name}: {str(e)}")
                continue
        
        if texts:
            progress_bar.progress(0.8, text="Creating vector database...")
            
            # Create text splitter
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
            
            # Split texts into chunks so in the textspiter the lanngchainDocument does in
            docs = [LangchainDocument(page_content=text) for text in texts]
            split_docs = text_splitter.split_documents(docs)
            
            # Get vector database and add documents
            vectordb = get_vectordb(db_name, embedding_manager.embeddings)
            
            # Use the appropriate method to add documents
            if hasattr(vectordb, 'add_documents'):
                vectordb.add_documents(split_docs)
                #this is working
                # print("Documents added to vector database thorugh the add_docuemnts")
            elif hasattr(vectordb, 'from_documents'):
                vectordb = vectordb.from_documents(split_docs, embedding_manager.embeddings)
                # print("Documents added to vector database thorugh the from_documents")
            else:
                raise AttributeError("Vector database doesn't support adding documents")
            
            progress_bar.progress(1.0, text="Database created successfully!")
            st.success(f"Successfully created database '{db_name}' with {len(texts)} documents!")
            return True
        
        return False
        
    except Exception as e:
        st.error(f"Error creating database: {str(e)}")
        logging.error(f"Database creation error: {str(e)}")
        return False
    
# Database sucessfuly created till here --------------------------------------------------------------------------


#Not using this one
def extract_headings(doc: DocxDocument) -> List[str]:
    """Extract all headings from the document."""
    headings = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # Check for heading style
        if paragraph.style.name == "Heading":
            headings.append(text)
        # Check for numbered sections (e.g., "1. Introduction" or "1.1 Background")
        elif re.match(r'^(?:\d+\.)*\d+\s+[A-Z]', text):
            headings.append(text)
    return headings



def extract_sections(doc: DocxDocument) -> List[Tuple[str, str]]:
    """Extract sections and their content from a document."""
    sections = []
    current_section = None
    current_content = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Match sections with or without word count
        if re.match(SECTION_PATTERN, text) or text in ["Purpose", "Objectives"]:  # Added common section titles
            # Save previous section if exists
            if current_section:
                sections.append((
                    current_section,
                    '\n'.join(current_content).strip()
                ))
            
            # Start new section
            current_section = text
            current_content = []
        elif current_section:  # Only add content if we're in a section
            current_content.append(text)
    
    # Add final section
    if current_section:
        sections.append((
            current_section,
            '\n'.join(current_content).strip()
        ))
    
    return sections

def analyze_document_structure(doc_path: str) -> Dict[str, Any]:
    """Analyze document structure by extracting headings and sections."""
    try:
        doc = DocxDocument(doc_path)
        sections = []
        current_section = None
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this is a heading
            is_heading = (
                (hasattr(paragraph, 'style') and paragraph.style.name.startswith('Heading')) or
                re.match(r'^\d+\..*', text) or  # Numbered sections like "1. Introduction"
                text in ["Purpose", "Objectives", "Introduction", "Methodology", 
                        "Results", "Discussion", "Conclusion", "Background", 
                        "Literature Review", "Abstract"]  # Common section titles
            )
            
            if is_heading:
                # Save previous section if exists
                if current_section:
                    sections.append({
                        'title': current_section,
                        'content': '\n'.join(current_content).strip()
                    })
                    current_content = []
                
                current_section = text
            elif current_section:
                current_content.append(text)
        
        # Add final section
        if current_section and current_content:
            sections.append({
                'title': current_section,
                'content': '\n'.join(current_content).strip()
            })
        
        return {'sections': sections}
        
    except Exception as e:
        st.error(f"Error analyzing document structure: {str(e)}")
        logging.error(f"Document analysis error: {str(e)}\n{traceback.format_exc()}")
        return {'sections': []}

def extract_word_limit(text: str) -> Optional[int]:
    """Extract word limit from text if present."""
    match = re.search(r'\((\d+)\s*words?\)', text)
    return int(match.group(1)) if match else None

def extract_requirements(text: str) -> List[str]:
    """Extract requirements from section text."""
    requirements = []
    
    # Common requirement patterns
    if "must include" in text.lower():
        requirements.append("Must include specific elements")
    if "required" in text.lower():
        requirements.append("Required section")
    if re.search(r'\(.*?required.*?\)', text.lower()):
        requirements.append("Has specific requirements")
    
    return requirements

def update_table_content(doc: DocxDocument, table_index: int, row: int, col: int, value: str) -> DocxDocument:
    """Update content of a specific table cell."""
    try:
        if table_index < len(doc.tables):
            table = doc.tables[table_index]
            if row < len(table.rows) and col < len(table.rows[0].cells):
                table.cell(row, col).text = value
        return doc
    except Exception as e:
        st.error(f"Error updating table: {str(e)}")
        return doc

def get_section_content(section_title: str, vectordb, custom_prompt: str = None) -> str:
    """Generate content for a section using vector database."""
    try:
        # Create retriever from vectorstore
        retriever = vectordb.as_retriever(search_kwargs={"k": 3})
        
        # Get relevant documents
        relevant_docs = retriever.get_relevant_documents(section_title)
        context = "\n\n".join([doc.page_content for doc in relevant_docs])
        
        # Create generation prompt
        prompt = f"""
        Write content for section '{section_title}'.
        
        Use this reference material for context:
        {context}
        
        {f'Additional instructions: {custom_prompt}' if custom_prompt else ''}
        
        Guidelines:
        1. Write professionally and clearly
        2. Include specific examples and details
        3. Stay focused on the section topic
        4. Be concise but comprehensive
        5. Use natural transitions
        
        Write the content now:
        """
        
        # Get response from LLM
        response = llm_manager.reasoning_model.invoke(prompt)
        content = response.content if hasattr(response, 'content') else str(response)
        return content.strip()
        
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return ""

def get_complete_document() -> DocxDocument:
    """Get the complete document with all sections."""
    if 'document_analysis' not in st.session_state:
        raise ValueError("No document analysis found in session state")
    
    try:
        # Create new document
        doc = DocxDocument()
        
        # Get sections from document analysis
        sections = st.session_state.document_analysis.get('sections', [])
        
        # Add each section
        for section_title, original_content, _ in sections:
            # Get edited content from session state if it exists
            section_key = f"section_{section_title}"
            content = st.session_state.get(section_key, original_content)
            
            # Add section title
            doc.add_heading(section_title, level=1)
            
            # Add content
            if content:
                doc.add_paragraph(content)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        return doc
        
    except Exception as e:
        st.error(f"Error generating document: {str(e)}")
        raise

def update_section_content(doc: DocxDocument, section_title: str, new_content: str) -> DocxDocument:
    """Update content of a specific section in the document."""
    try:
        found_section = False
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == section_title:
                found_section = True
                # Delete existing content until next section
                j = i + 1
                while j < len(doc.paragraphs):
                    if is_heading(doc.paragraphs[j].text.strip()):
                        break
                    doc.paragraphs[j]._element.getparent().remove(doc.paragraphs[j]._element)
                
                # Add new content after the section title
                new_para = doc.add_paragraph(new_content)
                # Move the new paragraph after the section title
                paragraph._element.addnext(new_para._element)
                break
        
        if not found_section:
            logging.warning(f"Section '{section_title}' not found in document")
        
        return doc
    except Exception as e:
        error_msg = f"Error updating section content: {str(e)}"
        st.error(error_msg)
        logging.error(error_msg)
        return doc

def is_heading(text: str) -> bool:
    """Check if text is a heading."""
    return (
        re.match(r'^\d+\..*', text) or
        text in ["Purpose", "Objectives", "Introduction", "Methodology", 
                "Results", "Discussion", "Conclusion", "Background", 
                "Literature Review", "Abstract"]
    )

def on_section_update(section_title: str, new_content: str):
    """Callback function for section updates."""
    try:
        doc = DocxDocument(st.session_state.template_path)
        doc = update_section_content(doc, section_title, new_content)
        doc.save(st.session_state.template_path)
        return True
    except Exception as e:
        st.error(f"Error updating section: {str(e)}")
        return False

def fill_document_template(template_path: str, replacements: Dict[str, str]) -> DocxDocument:
    """Fill a Word document template with provided values.
    
    Args:
        template_path: Path to the template document
        replacements: Dictionary of placeholder-value pairs
        
    Returns:
        Filled Document object
    """
    doc = DocxDocument(template_path)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)
    
    return doc

def create_document_with_sections(sections: List[Tuple[str, str]], title: str = None) -> DocxDocument:
    """Create a new document with the given sections."""
    try:
        doc = DocxDocument()
        
        # Add title if provided
        if title:
            doc.add_heading(title, 0)
        
        # Add each section
        for section_title, content in sections:
            # Add section title
            doc.add_heading(section_title, level=1)
            
            # Add content
            if content:
                doc.add_paragraph(content)
            
            # Add spacing
            doc.add_paragraph()
        
        return doc
        
    except Exception as e:
        st.error(f"Error creating document: {str(e)}")
        raise

def save_document_docx(doc: DocxDocument, output_path: str):
    """Save a Document object to a file."""
    doc.save(output_path)

# Main application layout
left_col, middle_col, right_col = st.columns([1, 2, 1])

# Left column - Document Template and Database Management
with left_col:
    with st.container():
        st.markdown("### Document Management")
        
        # Database creation section
        with st.expander("Create New Database", expanded=False):
            new_db_name = st.text_input("Database Name", placeholder="Enter database name")
            uploaded_files = st.file_uploader(
                "Upload Documents",
                type=["txt", "docx"],
                accept_multiple_files=True,
                help="Upload research-related documents to create knowledge base"
            )
            
            if uploaded_files and new_db_name:
                if st.button("Create Database"):
                    if create_database(new_db_name, uploaded_files):
                        st.success(f"Database '{new_db_name}' created successfully!")
                        time.sleep(1)  # Give user time to see the success message
                        st.rerun()
        
        st.markdown("---")
        
        # Template upload and database selection
        uploaded_template = st.file_uploader(
            "Upload Document Template",
            type=SUPPORTED_FORMATS,
            help="Upload a document template"
        )
        
        if uploaded_template:
            try:
                # Create a temporary file for the uploaded document
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    tmp_file.write(uploaded_template.getvalue())
                    st.session_state.template_path = tmp_file.name
                    # Update global template path
                
                    _TEMPLATE_PATH = tmp_file.name
                
                # Analyze document structure
                analysis = analyze_document_structure(st.session_state.template_path)
                st.session_state.document_analysis = analysis
                st.session_state.sections = analysis['sections']
                
                # Display document structure
                st.success("Template uploaded and analyzed successfully!")
                
                # Show document structure
                st.markdown("#### Document Structure")
                
                # Display sections
                for section_title, content in analysis['sections']:
                    st.markdown(f"- {section_title}")
                
                st.write(f"Found {len(analysis['sections'])} total sections")
                
                # Database selection
                available_dbs = get_available_databases()
                if available_dbs:
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        selected_db = st.selectbox(
                            "Select Knowledge Base",
                            options=available_dbs,
                            key="db_selector"
                        )
                    
                    with col2:
                        selected_model = st.selectbox(
                            "Select Model",
                            options=QUERY_MODELS,
                            format_func=lambda x: f"🤖 {x}",
                            key="model_selector"
                        )
                    
                    if st.button("Load Database"):
                        with st.spinner("Loading database..."):
                            vectordb = get_vectordb(selected_db, embedding_manager.embeddings)
                            st.session_state.vectordb = vectordb
                            st.session_state.chain = llm_manager.create_chain(
                                vectordb,
                                model_name=selected_model  # Pass the selected model
                            )
                            st.success(f"Database '{selected_db}' loaded successfully!")
                
                # Document filling section
                if st.session_state.vectordb:
                    st.markdown("### Fill Document")
                    form_data = {}
                    
                    # # Basic document metadata
                    # form_data['title'] = st.text_input("Document Title", help="Enter the title of the document")
                    # form_data['date'] = st.text_input("Document Date", help="Enter date (YYYY-MM-DD)")
                    # form_data['version'] = st.text_input("Document Version", help="Enter version number (e.g., 1.0)")
                    # form_data['organization'] = st.text_input("Organization Name", help="Enter your organization name")
                    
                    if st.button("Fill Document", use_container_width=True):
                        try:
                            with st.spinner("Filling document..."):
                                # Start with the template
                                doc = DocxDocument(st.session_state.template_path)
                                
                                # Fill each section
                                for section in st.session_state.document_analysis['sections']:
                                    section_title = section.get('title', '')
                                    section_key = f"section_{section_title}"
                                    
                                    # Get existing content or generate new
                                    content = st.session_state.get(section_key)
                                    if not content:
                                        content = get_section_content(section_title, st.session_state.vectordb)
                                    
                                    if content:
                                        doc = update_section_content(doc, section_title, content)
                                        st.session_state[section_key] = content
                                
                                # Save document
                                output_path = "filled_document.docx"
                                doc.save(output_path)
                                
                                # Offer download
                                with open(output_path, "rb") as file:
                                    st.download_button(
                                        label="Download Filled Document",
                                        data=file,
                                        file_name="filled_document.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True
                                    )
                                
                                st.success("Document filled successfully!")
                                
                        except Exception as e:
                            st.error(f"Error filling document: {str(e)}")
                else:
                    st.warning("Please load a knowledge base before filling the document.")
            
            except Exception as e:
                st.error(f"Error processing template: {str(e)}")
                if st.session_state.template_path and os.path.exists(st.session_state.template_path):
                    os.unlink(st.session_state.template_path)
                # Reset session state on error
                st.session_state.template_path = None
                st.session_state.document_analysis = None
                st.session_state.sections = []
                st.session_state.current_section = None

# Middle column - Document Content
with middle_col:
    st.markdown("### Document Content")
    if st.session_state.document_analysis:
        sections = st.session_state.document_analysis.get('sections', [])
        if not sections:
            st.warning("No sections found in the document. Please make sure your document has proper headings or section titles.")
        else:
            for section in sections:
                section_title = section.get('title', '')
                with st.expander(section_title, expanded=True):
                    # Section content
                    section_key = f"section_{section_title}"
                    current_content = st.session_state.get(section_key, section.get('content', ''))
                    
                    # Content editing area
                    new_content = st.text_area(
                        "Content",
                        value=current_content,
                        height=200,
                        key=f"display_{section_key}"
                    )
                    
                    # Generation instructions
                    custom_prompt = st.text_area(
                        "Instructions for AI",
                        key=f"prompt_{section_key}",
                        height=100,
                        placeholder="Give instructions to AI for generating content..."
                    )
                    
                    # Action buttons
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Generate", key=f"gen_{section_key}"):
                            if st.session_state.vectordb:
                                with st.spinner("Generating content..."):
                                    prompt_to_use = custom_prompt.strip() if custom_prompt else None
                                    generated_content = get_section_content(
                                        section_title,
                                        st.session_state.vectordb,
                                        prompt_to_use
                                    )
                                    if generated_content:
                                        st.session_state[section_key] = generated_content
                                        on_section_update(section_title, generated_content)
                                        st.success("Content generated!")
                                        st.rerun()
                            else:
                                st.warning("Please load a knowledge base first")
                    
                    with col2:
                        if st.button("Clear", key=f"clear_{section_key}"):
                            st.session_state[section_key] = ""
                            on_section_update(section_title, "")
                            st.success("Content cleared!")
                            st.rerun()
                    
                    st.markdown("---")

# Right column - Document Actions
with right_col:
    st.markdown("### Document Actions")
    if st.session_state.document_analysis and st.session_state.document_analysis.get('sections'):
        if st.button("Generate & Download Document", use_container_width=True):
            try:
                # Create new document
                doc = DocxDocument()
                
                # Add sections with their content
                for section in st.session_state.document_analysis['sections']:
                    section_title = section.get('title', '')
                    section_key = f"section_{section_title}"
                    content = st.session_state.get(section_key, '')
                    
                    # Add section title
                    doc.add_paragraph(section_title).style = 'Heading 1'
                    
                    # Add section content
                    if content:
                        doc.add_paragraph(content)
                    
                    # Add spacing between sections
                    doc.add_paragraph()
                
                # Save to bytes
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                # Show download button
                st.download_button(
                    label="Download Document",
                    data=doc_bytes,
                    file_name="generated_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.success("Document generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating document: {str(e)}")
                logging.error(f"Document generation error: {str(e)}")
    else:
        st.warning("Please upload a document template first")

# Right column - Chat Assistant
with right_col:
    st.markdown("### Chat Assistant")
    
    st.info("""
    Ask me questions about the document sections. For example:
    - What should I include in this section?
    - How do I improve this content?
    - What are the best practices?
    """)
    
    chat_container = st.container()
    with chat_container:
        # Display messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    
    # Chat input
    if query := st.chat_input("Ask for assistance..."):
        if not st.session_state.chain:
            st.error("Please load a database first!")
        else:
            st.session_state.messages.append({"role": "user", "content": query})
            
            with st.spinner("Processing..."):
                result = asyncio.run(process_query_async(query))
                
                if result:
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": result["answer"]
                    })
            
            st.rerun()

# Add a footer
st.markdown("---")
st.markdown(
    "Made by Ashu"
    "Using vector databases for intelligent form filling and document editing"
)

class DocumentProcessor:
    """Handle document processing and management"""
    
    def __init__(self):
        self.thread_pool = ThreadPoolExecutor(max_workers=2)
    
    async def process_section(self, section_title: str, content: str) -> Optional[str]:
        """Process document sections asynchronously"""
        try:
            return await asyncio.get_event_loop().run_in_executor(
                self.thread_pool,
                lambda: self._process_section_content(section_title, content)
            )
        except Exception as e:
            st.error(f"Error processing section {section_title}: {str(e)}")
            return None
    
    def _process_section_content(self, title: str, content: str) -> str:
        # Add your section processing logic here
        return content

# Add proper cleanup
def cleanup_files():
    """Clean up temporary files safely"""
    try:
        if st.session_state.template_path and os.path.exists(st.session_state.template_path):
            os.unlink(st.session_state.template_path)
    except Exception as e:
        logging.error(f"Cleanup error: {str(e)}")

# Register cleanup
atexit.register(cleanup_files)

def enhance_section_content(section_title: str, current_content: str, custom_prompt: str, vectordb) -> str:
    """Enhance existing section content based on custom prompt.
    
    Args:
        section_title: Title of the section
        current_content: Current content of the section
        custom_prompt: User's custom instructions
        vectordb: Vector database for additional context
        
    Returns:
        Enhanced content string
    """
    try:
        # Get additional context from vector database
        retriever = vectordb.as_retriever(search_kwargs={"k": 3})
        relevant_docs = retriever.get_relevant_documents(section_title)
        additional_context = "\n\n".join([doc.page_content for doc in relevant_docs])
        
        # Create enhancement prompt
        enhancement_prompt = f"""
        Current content for section '{section_title}':
        
        {current_content}
        
        User's instructions:
        {custom_prompt}
        
        Additional reference material:
        {additional_context}
        
        Please rewrite the content following these guidelines:
        1. Maintain the core message and key points from the current content
        2. Incorporate the user's specific instructions
        3. Use relevant information from the reference material
        4. Ensure natural flow and professional tone
        5. Keep similar length to original unless specified otherwise
        
        Enhanced content:
        """
        
        # Get enhanced content
        response = llm_manager.reasoning_model.invoke(enhancement_prompt)
        enhanced_content = response.content if hasattr(response, 'content') else str(response)
        
        return enhanced_content.strip()
        
    except Exception as e:
        st.error(f"Error enhancing content: {str(e)}")
        return current_content

